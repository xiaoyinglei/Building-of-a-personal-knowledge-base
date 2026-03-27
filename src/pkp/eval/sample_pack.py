from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path

import fitz  # type: ignore[import-untyped]
from docx import Document as WordDocument
from PIL import Image, ImageDraw

from pkp.eval.models import OfflineEvalFixture, OfflineEvalQuestion
from pkp.repo.interfaces import OcrRegion, OcrResult, OcrVisionRepo
from pkp.types.content import SourceType


@dataclass(frozen=True)
class BuiltinEvalPack:
    fixtures: list[OfflineEvalFixture]
    questions: list[OfflineEvalQuestion]
    ocr_repo: OcrVisionRepo
    question_bank_path: Path


class BuiltinEvalOcrRepo:
    def extract(self, image_path: Path) -> OcrResult:
        if image_path.name != "dashboard-metrics.png":
            return OcrResult(visible_text="", visual_semantics="")
        return OcrResult(
            visible_text="运营仪表盘 图2 收入 1280 订单 342",
            visual_semantics="一个包含 KPI 卡片的运营仪表盘截图",
            regions=[
                OcrRegion(text="图2", bbox=(16, 16, 120, 52)),
                OcrRegion(text="收入 1280", bbox=(16, 64, 220, 120)),
                OcrRegion(text="订单 342", bbox=(16, 128, 220, 184)),
            ],
        )


def prepare_builtin_eval_pack(output_dir: Path) -> BuiltinEvalPack:
    output_dir.mkdir(parents=True, exist_ok=True)
    fixtures = [
        _write_markdown_fixture(output_dir / "quarterly-review.md"),
        _write_docx_fixture(output_dir / "operations-brief.docx"),
        _write_pdf_fixture(output_dir / "research-notes.pdf"),
        _write_image_fixture(output_dir / "dashboard-metrics.png"),
    ]
    question_bank_path = _repo_root() / "data" / "eval" / "questions.json"
    payload = json.loads(question_bank_path.read_text(encoding="utf-8"))
    questions = [OfflineEvalQuestion.model_validate(item) for item in payload["questions"]]
    return BuiltinEvalPack(
        fixtures=fixtures,
        questions=questions,
        ocr_repo=BuiltinEvalOcrRepo(),
        question_bank_path=question_bank_path,
    )


def _write_markdown_fixture(path: Path) -> OfflineEvalFixture:
    path.write_text(
        "# 季度复盘\n\n"
        "## KPI 表\n\n"
        "| 指标 | 数值 |\n"
        "| --- | ---: |\n"
        "| 收入 | 1280 |\n"
        "| 订单 | 342 |\n"
        "| 续约率 | 91% |\n\n"
        "## 收入分析\n\n"
        "收入增长主要来自两个原因。第一，企业续约增长了，"
        "客服把续约流程从五步压缩到两步后，签回周期更短。"
        "第二，自助升级增加了，计费页面改版后，"
        "试用用户更容易完成升级。财务团队补充说，"
        "这两项变化同时作用在老客户和新试点项目上。\n\n"
        "## 发布门槛\n\n"
        "上线评审不是只看单个指标，而是看一整段闭环证据。"
        "第一道门槛是离线评估包里的检索精度必须稳定在 95% 以上，"
        "而且连续两轮都不能回落。评审会还要求把 child 命中、"
        "table 命中和 parent 回填的结果分别记录下来，"
        "避免只看一个漂亮的总体分数。"
        "团队在这个阶段会反复抽样检查 chunk 有没有被切坏，"
        "尤其是长段落和多句解释型内容，因为这些地方最容易把答案拆碎。"
        "第二道门槛是 metadata 审计必须显示必填字段缺失为零，"
        "至少要覆盖 location、toc_path、chunk_role 和 content_hash。"
        "只有两个门槛同时满足，部署窗口才会打开。\n",
        encoding="utf-8",
    )
    return OfflineEvalFixture(
        fixture_id="quarterly_review_markdown",
        filename=path.name,
        source_type=SourceType.MARKDOWN,
        description=(
            "Markdown fixture with a table, a direct child-answer section, "
            "and a long parent-backfill section."
        ),
        path=path,
    )


def _write_docx_fixture(path: Path) -> OfflineEvalFixture:
    document = WordDocument()
    document.add_heading("运营周报", level=1)
    document.add_paragraph("本周值班目标是减少回归成本。")
    document.add_heading("值班指标", level=2)
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "待处理告警"
    table.cell(1, 1).text = "7"
    table.cell(2, 0).text = "自动修复"
    table.cell(2, 1).text = "5"
    document.add_heading("值班结论", level=2)
    document.add_paragraph(
        "运营组把凌晨告警清理时间从 42 分钟压缩到了 18 分钟，"
        "主要靠统一告警标签和更短的回滚手册。"
    )
    document.save(str(path))
    return OfflineEvalFixture(
        fixture_id="operations_brief_docx",
        filename=path.name,
        source_type=SourceType.DOCX,
        description="DOCX fixture with headings, a table, and one direct answer paragraph.",
        path=path,
    )


def _write_pdf_fixture(path: Path) -> OfflineEvalFixture:
    document = fitz.open()
    first_page = document.new_page()
    first_page.insert_text(
        (72, 72),
        "Research Notes\nFast Path should answer direct questions with citations.",
    )
    second_page = document.new_page()
    second_page.insert_text(
        (72, 72),
        (
            "Deep Path should decompose research questions, expose conflicts, "
            "and report uncertainty.\n"
            "Risk review says supply chain delay remains manageable when "
            "fallback suppliers are active."
        ),
    )
    document.save(path)
    document.close()
    return OfflineEvalFixture(
        fixture_id="research_notes_pdf",
        filename=path.name,
        source_type=SourceType.PDF,
        description="PDF fixture with two pages and direct retrieval statements.",
        path=path,
    )


def _write_image_fixture(path: Path) -> OfflineEvalFixture:
    image = Image.new("RGB", (520, 240), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((18, 18), "Figure 2", fill="black")
    draw.text((18, 82), "Revenue 1280", fill="black")
    draw.text((18, 146), "Orders 342", fill="black")
    image.save(path)
    return OfflineEvalFixture(
        fixture_id="dashboard_metrics_image",
        filename=path.name,
        source_type=SourceType.IMAGE,
        description="Image fixture for OCR-region and image-summary evaluation.",
        path=path,
    )


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[3]
